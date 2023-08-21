"""
assignmentTwo.py
The Assignment 2 of COMP1112-01-235

Developer : Dain Shin
Date: July 11, 2023
"""

# Importing docx for word and openpyxl for excel
from docx import Document
from openpyxl import Workbook

# Creating Workbook and activate the worksheet in the excel file
wb = Workbook()
ws = wb.active

# Writing the header in the excel file
ws.append(['Invoice Number', 'Total Quantity', 'Subtotal', 'Tax', 'Total'])

# Using for loop, open the existing word file named INV1000000 ~ INV1000200
for i in range(1000000, 1000200):

        existingDoc = Document("INV{}.docx".format(i))

        # There are 3 parsts in the word document: Heading(Invoice number), paragraph1(Product information) and paragraph2(Subtotal, tax and total)
        # Extracting the invoice number from the first paragraph
        invoiceNum = existingDoc.paragraphs[0].text
 
        # Reading the text line by line and putting them into paraLine list
        para = existingDoc.paragraphs[1].text
        paraLine = para.strip().split("\n")  

        # In products dictionary, the name of the product and quantity will be stored as a pair of key and value
        # sumProductNum will store the total quantity of the products
        products = {}
        sumProductNum = 0
        for item in paraLine[1:]:
                key, value = item.split(":")
                products[key] = value
                sumProductNum += int(value)

        # Reading the text line by line and putting them into paraLine2 list
        para2 = existingDoc.paragraphs[2].text
        paraLine2 = para2.strip().split("\n")


        # In prices dictionary, the data will be stored as a paif of key and value
        prices = {}
        for i in paraLine2:
                key, value = i.split(":")
                prices[key] = float(value)  # I changed the value into float type to make it numeric value in the excel file
                

        # Making a list for the needed data        
        list = [invoiceNum, sumProductNum, prices["SUBTOTAL"], prices["TAX"], prices["TOTAL"]]

        # Wrting the extracted data in the excel file
        ws.append(list)

        # Save the excel file
        wb.save(filename="assingmentTwo.xlsx")